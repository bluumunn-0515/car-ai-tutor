# -*- coding: utf-8 -*-
"""구두보고용_설명문.md → 구두보고용_설명문.docx 변환"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from md_to_docx import (
    BODY_SIZE,
    EN_FONT,
    H2_SIZE,
    H3_SIZE,
    KO_FONT,
    add_markdown_table,
    add_para,
    is_table_separator,
    parse_inline,
    set_run_font,
    split_table_row,
)


SRC = Path(__file__).resolve().parent / "구두보고용_설명문.md"
DST = Path(__file__).resolve().parent / "구두보고용_설명문.docx"


def add_title(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(18)
    run = p.add_run(text)
    set_run_font(run, size_pt=18, bold=True)


def add_h2(doc, text: str):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, size_pt=H2_SIZE, bold=True)


def add_h3(doc, text: str):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.4)
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, size_pt=H3_SIZE, bold=True)


def add_quote(doc, text: str):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.8)
    pf.right_indent = Cm(0.5)
    pf.first_line_indent = Cm(0)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size_pt=BODY_SIZE, color=RGBColor(0x1A, 0x1A, 0x1A))


def add_bullet(doc, text: str, level: int = 0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.8 + level * 0.4)
    pf.first_line_indent = Cm(-0.35)
    pf.space_after = Pt(3)
    run = p.add_run("• " + text)
    set_run_font(run, size_pt=BODY_SIZE)


def add_numbered(doc, text: str):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.8)
    pf.first_line_indent = Cm(-0.35)
    pf.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size_pt=BODY_SIZE)


def convert(md_path: Path, dst_path: Path):
    md_text = md_path.read_text(encoding="utf-8")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(BODY_SIZE)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), KO_FONT)

    lines = md_text.splitlines()
    i = 0
    first_h1 = True

    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        if re.match(r"^\s*---+\s*$", line):
            i += 1
            continue

        if "|" in line and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            header = split_table_row(line)
            i += 2
            body = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                body.append(split_table_row(lines[i]))
                i += 1
            add_markdown_table(doc, header, body)
            doc.add_paragraph()
            continue

        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            text = parse_inline(m.group(2)).strip()
            if level == 1:
                if first_h1:
                    add_title(doc, text)
                    first_h1 = False
                else:
                    add_h2(doc, text)
            elif level == 2:
                add_h2(doc, text)
            else:
                add_h3(doc, text)
            i += 1
            continue

        if line.strip().startswith(">"):
            text = parse_inline(line.strip().lstrip(">").strip())
            add_quote(doc, text)
            i += 1
            continue

        if re.match(r"^-\s+", line.strip()):
            text = parse_inline(line.strip()[2:]).strip()
            add_bullet(doc, text)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line.strip()):
            text = parse_inline(line.strip()).strip()
            add_numbered(doc, text)
            i += 1
            continue

        if line.strip().startswith("**") and line.strip().endswith("**"):
            text = parse_inline(line.strip())
            add_para(doc, text, bold=True, first_indent_cm=0)
            i += 1
            continue

        text = parse_inline(line.strip())
        if text.startswith("→"):
            add_para(doc, text, first_indent_cm=0.5, left_indent_cm=0.3)
        else:
            add_para(doc, text)
        i += 1

    doc.save(dst_path)
    print(f"[OK] saved: {dst_path}")
    print(f"     size : {dst_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    convert(SRC, DST)
