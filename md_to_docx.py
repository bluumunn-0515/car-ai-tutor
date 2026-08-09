"""
연구계획서_수정본_v2.md  →  연구계획서_수정본_v2.docx

한국 학위논문 형식에 적합한 DOCX 변환기.
- 맑은 고딕 / 함초롬바탕 폰트 (한글 워드프로세서가 자동 매핑)
- Ⅰ, 1, 가, 1), 가), (1) 위계의 들여쓰기·글자 크기 자동 적용
- 마크다운 표(`|`) → 정식 표 변환
- 코드블록 → 고정폭 폰트 박스
- 한컴오피스(한글)에서 .docx 파일을 열고 [파일 > 다른 이름으로 저장 > .hwp] 로 저장 가능
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement


# ----- 설정 -----
SRC = Path(__file__).resolve().parent / "연구계획서_수정본_v3.md"
DST = Path(__file__).resolve().parent / "연구계획서_수정본_v3.docx"
FIGURES_DIR = Path(__file__).resolve().parent / "figures"

# 앱 실제 캡처 권장 ([그림 12]~[그림 22])
APP_SCREENSHOT_FIGS = set(range(12, 23))

KO_FONT = "맑은 고딕"
EN_FONT = "Times New Roman"
MONO_FONT = "D2Coding"  # 없으면 한글이 Consolas로 폴백

BODY_SIZE = 11  # pt
H1_SIZE = 18    # Ⅰ ~ Ⅴ 장 제목
H2_SIZE = 15    # 1., 2., 3., …
H3_SIZE = 13    # 가., 나., …
H4_SIZE = 12    # 1), 2), …
H5_SIZE = 12    # 가), 나), (1), …


def set_run_font(run, ko=KO_FONT, en=EN_FONT, size_pt=BODY_SIZE,
                 bold=False, color=None):
    """한·영 폰트를 모두 지정 (한컴오피스에서 한글이 깨지지 않도록)."""
    run.font.name = en
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), ko)
    rFonts.set(qn("w:ascii"), en)
    rFonts.set(qn("w:hAnsi"), en)
    rFonts.set(qn("w:cs"), en)


def add_para(doc, text, size_pt=BODY_SIZE, bold=False, align=None,
             first_indent_cm=0.7, left_indent_cm=0.0, space_after=4,
             color=None, ko=KO_FONT):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.first_line_indent = Cm(first_indent_cm)
    pf.left_indent = Cm(left_indent_cm)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.6
    run = p.add_run(text)
    set_run_font(run, ko=ko, size_pt=size_pt, bold=bold, color=color)
    return p


def add_heading_para(doc, text, level):
    """Ⅰ, 1, 가, 1), 가) 위계별 스타일."""
    if level == 1:
        # 새 페이지에서 시작
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size_pt=H1_SIZE, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(18)
        return
    sizes = {2: H2_SIZE, 3: H3_SIZE, 4: H4_SIZE, 5: H5_SIZE, 6: H5_SIZE}
    indents = {2: 0.0, 3: 0.5, 4: 1.0, 5: 1.5, 6: 2.0}
    size_pt = sizes.get(level, BODY_SIZE)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(indents.get(level, 0.0))
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(10 if level == 2 else 6)
    pf.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size_pt=size_pt, bold=True)


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.8)
    pf.first_line_indent = Cm(0)
    pf.space_after = Pt(8)
    run = p.add_run(code_text)
    set_run_font(run, ko=MONO_FONT, en=MONO_FONT, size_pt=10,
                 color=RGBColor(0x33, 0x33, 0x33))


def set_cell_borders(cell):
    """표 셀에 얇은 검은 테두리."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "000000")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_markdown_table(doc, header_cells, body_rows):
    n_cols = len(header_cells)
    table = doc.add_table(rows=1 + len(body_rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 헤더
    hdr = table.rows[0].cells
    for i, h in enumerate(header_cells):
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        set_run_font(run, size_pt=BODY_SIZE, bold=True)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_borders(hdr[i])
        # 헤더 배경색
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "E8EEF7")
        tcPr.append(shd)

    # 본문
    for r_idx, row_data in enumerate(body_rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx in range(n_cols):
            cells[c_idx].text = ""
            text = row_data[c_idx] if c_idx < len(row_data) else ""
            para = cells[c_idx].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = Cm(0)
            run = para.add_run(text)
            set_run_font(run, size_pt=10)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cells[c_idx])


def parse_inline(line):
    """**bold**, *em*, `code` 마크다운을 그대로 두고 단순 텍스트로 반환.
    추후 필요시 run 단위 처리로 확장 가능. 여기서는 표식 제거만 수행."""
    s = line
    s = re.sub(r"`([^`]+)`", r"\1", s)
    s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
    s = re.sub(r"\*([^*]+)\*", r"\1", s)
    return s


def is_table_separator(line):
    return bool(re.match(r"^\s*\|?[\s\-:|]+\|[\s\-:|]+\|?\s*$", line))


def split_table_row(line):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def is_tab_table_row(line):
    """탭으로 구분된 표 행인지 판정 (탭이 1개 이상이면서 셀이 2개 이상)."""
    if not line.strip():
        return False
    return line.count("\t") >= 1


def split_tab_row(line):
    return [c.strip() for c in line.split("\t")]


def parse_figure_num(text: str) -> int | None:
    m = re.match(r"^\[그림\s+(\d+)\]", text.strip())
    return int(m.group(1)) if m else None


def figure_image_path(num: int) -> Path | None:
    path = FIGURES_DIR / f"fig_{num:02d}.png"
    return path if path.exists() else None


def add_figure_image(doc, path: Path, width_cm: float = 14.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))


def add_screenshot_placeholder(doc, caption: str):
    """앱 UI 스크린샷 삽입 위치 — 점선 박스 + 안내 문구."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)

    # 테두리 박스 (단락 음영 + 여백으로 시각적 구분)
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)

    run = p.add_run(
        "【 앱 스크린샷 삽입 예정 】\n"
        f"{caption}\n"
        "(앱 실행 화면을 캡처하여 이 위치에 이미지를 삽입하세요)"
    )
    set_run_font(run, size_pt=10, color=RGBColor(0x75, 0x75, 0x75))
    pf.left_indent = Cm(1.5)
    pf.right_indent = Cm(1.5)


def md_to_docx(md_text, dst_path):
    doc = Document()

    # 기본 여백 (한국 학위논문 표준: 위·아래 2.5cm, 좌·우 2.5cm)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # 기본 스타일 폰트
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(BODY_SIZE)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), KO_FONT)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i]

        # ----- 코드블록 -----
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                add_code_block(doc, "\n".join(code_buf))
                in_code = False
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ----- 빈 줄 -----
        if not line.strip():
            i += 1
            continue

        # ----- 수평 구분선 -----
        if re.match(r"^\s*---+\s*$", line):
            # 새 페이지로 처리
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)
            i += 1
            continue

        # ----- 표 (파이프 구분) -----
        if "|" in line and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            header = split_table_row(line)
            i += 2
            body = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                body.append(split_table_row(lines[i]))
                i += 1
            add_markdown_table(doc, header, body)
            continue

        # ----- 표 (탭 구분) -----
        if is_tab_table_row(line):
            header = split_tab_row(line)
            n_cols = len(header)
            j = i + 1
            body = []
            while j < len(lines) and is_tab_table_row(lines[j]):
                row = split_tab_row(lines[j])
                # 컬럼 수가 같거나 1개 정도만 차이나는 경우 본문 행으로 처리
                while len(row) < n_cols:
                    row.append("")
                if len(row) > n_cols:
                    # 마지막 셀에 나머지 합치기 (안전 처리)
                    row = row[: n_cols - 1] + ["\t".join(row[n_cols - 1 :])]
                body.append(row)
                j += 1
            # 행이 2줄(헤더 + 본문 1줄) 이상일 때만 표로 처리, 아니면 일반 본문
            if body:
                add_markdown_table(doc, header, body)
                i = j
                continue

        # ----- 헤더 -----
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            text = parse_inline(m.group(2)).strip()
            add_heading_para(doc, text, level)
            i += 1
            continue

        # ----- 본문 -----
        text = parse_inline(line).strip()
        # 위계가 명시된 한국식 번호 (가., 나., 1), 가) 등) 들여쓰기 보정
        if re.match(r"^[가-힣]\.\s", text):
            add_para(doc, text, first_indent_cm=0.5, left_indent_cm=0.5)
        elif re.match(r"^\d+\)\s", text):
            add_para(doc, text, first_indent_cm=0.5, left_indent_cm=1.0)
        elif re.match(r"^[가-힣]\)\s", text):
            add_para(doc, text, first_indent_cm=0.5, left_indent_cm=1.5)
        elif re.match(r"^\(\d+\)\s", text) or re.match(r"^[①-⑳]\s", text):
            add_para(doc, text, first_indent_cm=0.5, left_indent_cm=2.0)
        elif text.startswith("[그림"):
            fig_num = parse_figure_num(text)
            add_para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER,
                     first_indent_cm=0, space_after=4, bold=True)
            if fig_num is not None:
                img = figure_image_path(fig_num)
                if img is not None:
                    add_figure_image(doc, img)
                elif fig_num in APP_SCREENSHOT_FIGS:
                    add_screenshot_placeholder(doc, text)
            i += 1
            continue
        elif text.startswith("(편집 시 본 그림"):
            # 자동 생성 그림으로 대체되므로 편집 안내 문구는 생략
            i += 1
            continue
        elif text.startswith("<표"):
            add_para(doc, text, first_indent_cm=0, space_after=4, bold=True)
        else:
            add_para(doc, text)

        i += 1

    doc.save(dst_path)
    print(f"[OK] saved: {dst_path}")
    print(f"     size : {dst_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    import generate_figures
    generate_figures.main()

    md_text = SRC.read_text(encoding="utf-8")
    md_to_docx(md_text, DST)

    inserted = sum(1 for n in range(1, 29) if figure_image_path(n))
    placeholders = len(APP_SCREENSHOT_FIGS)
    print()
    print(f"[FIGURES] inserted: {inserted}, screenshot placeholders: {placeholders}")
    print()
    print("[HOW TO MAKE .HWP]")
    print("  1. Open Hancom Hangul Word Processor")
    print("  2. File > Open > select the .docx above")
    print("  3. File > Save As > file type [Hangul Document (*.hwp)]")
    print()
    print("[SCREENSHOTS] Insert real app captures at fig_12 ~ fig_22 placeholder boxes")
