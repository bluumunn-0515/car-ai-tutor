"""
연구계획서 표 1~7 → DOCX / JPG 개별 파일 생성.

출력:
  tables/표_01_NCS_능력단위.docx / .jpg
  ...
  tables/표_전체_1-7.docx  (7개 표 통합)
"""
from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path(__file__).parent / "tables"
DOWNLOAD_DIR = OUT_DIR / "download"
KO_FONT = "맑은 고딕"
EN_FONT = "Times New Roman"
BODY_SIZE = 11


# ── 표 데이터 (목차 번호·제목 기준) ────────────────────────────────
TABLES: list[dict] = [
    {
        "num": 1,
        "title": "NCS 자동차 정비 분야 세분류 15060303 능력단위 구성",
        "headers": ["번호", "능력단위 코드", "능력단위명"],
        "rows": [
            ["1", "LM1506030101", "자동차 충전장치 정비"],
            ["2", "LM1506030102", "자동차 시동장치 정비"],
            ["3", "LM1506030103", "자동차 냉·난방장치 정비"],
            ["4", "LM1506030104", "자동차 전기·전자회로 분석"],
            ["5", "LM1506030105", "자동차 편의장치 정비"],
            ["6", "LM1506030106", "자동차 등화장치 정비"],
            ["7", "LM1506030107", "자동차 주행안전장치 정비"],
            ["8", "LM1506030108", "자동차 네트워크통신장치 정비"],
            ["9", "LM1506030109", "하이브리드자동차 특화시스템 정비"],
            ["10", "LM1506030117", "전기자동차 특화시스템 정비"],
            ["11", "LM1506030118", "자동차 고전압전기장치 정비"],
        ],
        "col_widths_cm": [2.0, 4.5, 8.0],
        "merge_cols": None,
    },
    {
        "num": 2,
        "title": "Hattie & Timperley(2007) 피드백 3 질문과 4 수준",
        "headers": ["피드백 질문 / 수준", "본 프로젝트에서의 구현"],
        "rows": [
            ["Feed Up (Where am I going?)",
             "4단계 미션 카드의 단계 제목 + '✋ 생각해볼 점' 1줄 질문"],
            ["Feed Back (How am I going?)",
             "카테고리별 평가의 '사실' 행 + 'NCS 기준' 행 + '통과 / 보완' 라벨"],
            ["Feed Forward (Where to next?)",
             "카테고리별 평가의 '보완 제안' 행 + AI 찬스의 '✅ 정상이면 / ⚠ 의심되면' 분기"],
            ["과제 수준 (FT)", "측정값 단위·자릿수의 정/오 확인"],
            ["과정 수준 (FP)", "NCS 수행준거 절차와 학습자 수행 절차의 비교"],
            ["자기 조절 수준 (FR)",
             "'다시 평가 받기' 흐름 + 직전 평가 노출 + '오늘의 실습 소감' 입력"],
            ["자기 수준 (FS)",
             "의도적으로 차단 — 시스템 프롬프트에 '이모지·꾸밈문자 사용 금지', "
             "'잘했어요·최우수 등 단정적 평가어 금지' 규칙 명시"],
        ],
        "col_widths_cm": [5.5, 9.5],
        "merge_cols": None,
    },
    {
        "num": 3,
        "title": "Wood, Bruner & Ross(1976) 스캐폴딩 6 기능",
        "headers": ["번호", "기능", "설명"],
        "rows": [
            ["①", "관심 유발(Recruitment)", "학습자가 과제에 흥미를 갖고 참여하도록 유도"],
            ["②", "자유도 축소(Reduction in degrees of freedom)",
             "한 번에 다루어야 할 인지적 요소의 수를 제한"],
            ["③", "방향 유지(Direction maintenance)",
             "학습자가 목표를 잃지 않고 과제를 향해 나아가도록 유지"],
            ["④", "핵심 특성 표시(Marking critical features)",
             "과제에서 결정적인 요소를 부각하여 학습자의 주의 집중"],
            ["⑤", "좌절 통제(Frustration control)",
             "학습자가 좌절하지 않도록 적절한 수준의 도움 제공"],
            ["⑥", "시범(Demonstration)",
             "모범 답안의 형태를 제시하여 학습자가 그 형태를 모방·내재화"],
        ],
        "col_widths_cm": [1.8, 5.5, 7.7],
        "merge_cols": None,
    },
    {
        "num": 4,
        "title": "본 프로젝트의 NCS 능력단위별 평가 루브릭",
        "headers": ["NCS 능력단위", "카테고리 1", "카테고리 2", "카테고리 3", "카테고리 4"],
        "rows": [
            ["자동차 전기전자장치 고장진단", "안전·전원 차단 확인", "회로도/기호 분석",
             "회로시험기 측정 절차", "진단장비(스캐너) 활용"],
            ["배터리 점검", "배터리 외관/상태 확인", "개방회로 전압(OCV) 측정",
             "부하/CCA·SOC 판정", "암전류/배터리 센서 점검"],
            ["시동·충전장치 점검", "시동회로 점검", "발전기 출력 점검",
             "회로 전압강하 측정", "점검 절차/예비점검"],
            ["조명장치 점검", "등화회로 분석", "광원/전구 점검",
             "회로 전압/접지 측정", "BCM/CAN 등화 제어"],
            ["편의장치 점검", "편의장치 유형/회로 식별", "모듈 전원·접지 점검",
             "액추에이터/릴레이 점검", "스캐너 자기진단/강제구동"],
            ["네트워크 장치 점검", "통신 프로토콜 이해", "종단저항/배선 점검",
             "통신 신호/파형 측정", "게이트웨이/모듈 진단"],
        ],
        "col_widths_cm": [4.0, 3.0, 3.0, 3.0, 3.0],
        "merge_cols": None,
    },
    {
        "num": 5,
        "title": "프로젝트 개발 도구 정리",
        "headers": ["개발 단계/요소", "사용 도구/프로그램", "세부 역할/기능"],
        "rows": [
            ["기술 명세서 작성 및 개발 계획 정리", "Google Docs, 한글(HWP), Canva",
             "지속적인 개발 계획 업데이트 및 정리, 구현 화면 예시 참고용 이미지 제작"],
            ["개발 환경 (IDE)", "Cursor IDE, Python 3",
             "AI 기반 페어 프로그래밍, 코드 편집·수정·디버깅"],
            ["화면 구성 및 스타일 구현", "Streamlit 컴포넌트 + HTML/CSS 인라인",
             "반응형 화면 구성, 모바일·태블릿 호환 디자인"],
            ["멀티모달 입출력 처리",
             "Streamlit file_uploader, text_area, image + Pillow",
             "텍스트·이미지 입력, 단계별 사진 썸네일 압축 (Google Sheets 셀 50,000자 한도 적합화)"],
            ["AI 추론 엔진",
             "Google Gemini 2.5 Flash API (google-genai SDK) + 폴백",
             "단일 멀티모달 추론으로 텍스트+사진 통합 분석, 소크라테스식 발문 생성"],
            ["NCS 도메인 지식 통합", "NCS_RUBRIC 코드 사전 + 시스템 프롬프트",
             "NCS 능력단위별 하위 수행준거 키워드의 매 호출 자동 주입 (PeDK)"],
            ["피드백 생성",
             "시스템 프롬프트 + Hattie & Timperley(2007) 3질문×4수준 모형",
             "카테고리별 '사실→NCS 기준→보완 제안' 3행 구조의 평가, 한줄 요약, 종합 코멘트"],
            ["클라우드 데이터베이스",
             "Google Sheets (streamlit-gsheets-connection) + 60초 TTL 캐시",
             "users/history/final_assessments 3개 시트, read→merge→update 동시 편집 안전 패턴"],
            ["학습 활동 분석 및 시각화", "Plotly (방사형/막대 차트)",
             "NCS 단원별 평균, 카테고리별 분포 시각화 대시보드"],
            ["PDF 포트폴리오 생성", "fpdf2 + 한글 폰트 임베딩(malgun.ttf)",
             "학기 누적 기록의 단일 PDF 출력 (이미지 자동 압축 포함)"],
            ["코드 작성 및 수정·보완", "Cursor IDE 내장 Claude 3.5 Sonnet, GPT-4",
             "AI 기반 페어 프로그래밍, 리팩토링"],
            ["버전 관리", "Git, GitHub", "코드 버전 관리, 백업"],
            ["웹 배포", "Streamlit Community Cloud",
             "무료·공개 웹 배포, 학교 환경 무설치 접속"],
        ],
        "col_widths_cm": [4.0, 5.0, 7.0],
        "merge_cols": None,
    },
    {
        "num": 6,
        "title": "사전·사후 형성평가 루브릭",
        "headers": ["평가 영역", "평가 기준 (우수 5점 / 보통 3점 / 미흡 1점)"],
        "rows": [
            ["① 고장 증상 파악 및 가설 설정",
             "증상 인식의 정확성, 가설 범위의 구체성, 안전 절차 인식 여부"],
            ["② 회로도 분석 및 점검 계획 수립",
             "회로도 해석의 정확성, 전원→퓨즈→스위치→부하→접지 흐름 추적, 점검 우선순위 설정"],
            ["③ 계측 위치 선정 및 데이터 측정 절차",
             "측정 위치의 적절성, 멀티미터 모드·레인지 선정, 정상 기준값과의 비교 기준 명시"],
            ["④ 원인 도출 및 정비 전략 수립",
             "측정값 해석의 논리성, 원인 도출의 인과적 타당성, 다음 정비 조치의 구체성"],
        ],
        "col_widths_cm": [5.0, 10.0],
        "merge_cols": None,
    },
    {
        "num": 7,
        "title": "자동차 고장진단 앱에 대한 교사 인터뷰 질문",
        "headers": ["영역", "질문"],
        "rows": [
            ["인식", "1-1. 본 앱을 살펴보신 후, 전반적으로 어떤 인상이나 느낌을 받으셨나요?"],
            ["", "1-2. 이 앱이 기존 '자동차 전기·전자 제어' 실습 수업에서 보완하거나 "
             "새롭게 기여할 수 있는 부분이 있다면 무엇이라고 생각하시나요?"],
            ["기능 구성",
             "2-1. 앱의 기능 중에서(예: 4단계 미션 카드, AI 찬스 메커니즘, NCS 루브릭 평가, "
             "교사 모드 세특 자동 생성) 교육적으로 가장 유용하다고 생각되는 부분은 무엇인가요?"],
            ["", "2-2. 반대로, 개선이 필요하거나 학생들이 사용하기 어려울 것으로 "
             "예상되는 기능이 있다면 말씀해 주세요."],
            ["수업 적용 가능성과 기대 효과",
             "3-1. 실제 수업에 적용한다면, 이 앱을 어떤 방식이나 시점에서 활용할 수 있을 것 같나요? "
             "(예: 단원 도입, 실습 중 보조 도구, 수행평가 도구)"],
            ["", "3-2. 이 앱이 학생의 자동차 고장진단 능력 향상이나 자기 주도적 학습 태도에 "
             "미칠 긍정적 영향은 무엇이라고 예상하시나요? 또한 'AI 찬스 감점 메커니즘'이 "
             "학생의 AI 의존 행동에 어떤 영향을 미칠 것으로 보시는지 말씀해 주세요."],
        ],
        "col_widths_cm": [4.0, 11.0],
        "merge_cols": {0: [(0, 1), (2, 3), (4, 5)]},  # 영역 열 세로 병합
    },
]


def set_run_font(run, ko=KO_FONT, en=EN_FONT, size_pt=BODY_SIZE, bold=False):
    run.font.name = en
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), ko)
    rFonts.set(qn("w:ascii"), en)
    rFonts.set(qn("w:hAnsi"), en)


def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "000000")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def shade_cell(cell, fill="E8EEF7"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_table_to_doc(doc, spec: dict):
    """표 제목 + 본문 표를 Document에 추가."""
    num = spec["num"]
    title = spec["title"]
    headers = spec["headers"]
    rows = spec["rows"]
    col_widths = spec.get("col_widths_cm")
    merge_cols = spec.get("merge_cols") or {}

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(8)
    run = p.add_run(f"<표 {num}> {title}")
    set_run_font(run, size_pt=12, bold=True)

    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        set_run_font(run, bold=True)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_borders(cell)
        shade_cell(cell)

    # 본문
    for r_idx, row_data in enumerate(rows, start=1):
        for c_idx in range(n_cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            text = row_data[c_idx] if c_idx < len(row_data) else ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            run = para.add_run(text)
            set_run_font(run, size_pt=10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)

    # 세로 병합
    for col_idx, ranges in merge_cols.items():
        for start, end in ranges:
            top = table.rows[start + 1].cells[col_idx]
            bottom = table.rows[end + 1].cells[col_idx]
            top.merge(bottom)
            top.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # 표 아래 여백


def setup_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(BODY_SIZE)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), KO_FONT)
    return doc


def slug(num: int, title: str) -> str:
    short = {
        1: "NCS_능력단위",
        2: "Hattie_피드백",
        3: "Wood_스캐폴딩",
        4: "NCS_루브릭",
        5: "개발도구",
        6: "형성평가_루브릭",
        7: "교사_인터뷰",
    }
    return f"표_{num:02d}_{short[num]}"


def export_docx(spec: dict, path: Path):
    doc = setup_doc()
    add_table_to_doc(doc, spec)
    doc.save(path)


def _sanitize_for_jpg(text: str) -> str:
    """JPG 렌더링 시 맑은 고딕 미지원 이모지를 텍스트로 대체."""
    return (
        text.replace("✋", "[생각해볼 점]")
        .replace("✅", "[정상]")
        .replace("⚠", "[의심]")
    )


def _wrap_text(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    words, lines, cur = text.replace("\n", " "), [], ""
    # 한글은 글자 단위 줄바꿈
    for ch in words:
        if len(cur) >= max_chars and ch == " ":
            lines.append(cur)
            cur = ""
        cur += ch
        if len(cur) >= max_chars:
            lines.append(cur)
            cur = ""
    if cur:
        lines.append(cur)
    return "\n".join(lines)


def export_rubric_image(spec: dict, path: Path):
    """표 4 NCS 루브릭 — 연구계획서 삽입용 고해상도 PNG."""
    plt.rcParams.update({
        "font.family": "Malgun Gothic",
        "axes.unicode_minus": False,
    })

    headers = spec["headers"]
    rows = spec["rows"]
    n_cols = len(headers)

    fig, ax = plt.subplots(figsize=(16, 7.4))
    ax.axis("off")
    fig.suptitle(
        "<표 5> 본 프로젝트의 NCS 능력단위별 평가 루브릭",
        fontsize=15,
        fontweight="bold",
        y=0.98,
        color="#1A237E",
    )

    cell_text = [[str(c) for c in row] for row in rows]
    col_labels = list(headers)

    col_widths = [0.26, 0.185, 0.185, 0.185, 0.185]
    tbl = ax.table(
        cellText=cell_text,
        colLabels=col_labels,
        loc="center",
        cellLoc="center",
        colWidths=col_widths,
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(12)
    tbl.scale(1.0, 2.35)

    header_fc = "#2E5090"
    header_tc = "#FFFFFF"
    row_fc = ["#FFFFFF", "#F7F9FC"]
    unit_fc = ["#EEF2F7", "#E8EEF5"]
    cat_header_fc = ["#D5E8D4", "#BBDEFB", "#FFE0B2", "#FFCDD2"]

    for (row, col), cell in tbl.get_celld().items():
        cell.set_edgecolor("#B0BEC5")
        cell.set_linewidth(1.2)
        cell.PAD = 0.14
        if row == 0:
            if col == 0:
                cell.set_facecolor(header_fc)
            else:
                cell.set_facecolor(cat_header_fc[col - 1])
            cell.set_text_props(fontweight="bold", color=header_fc if col == 0 else "#263238", ha="center")
            if col == 0:
                cell.get_text().set_color(header_tc)
        else:
            data_row = row - 1
            if col == 0:
                cell.set_facecolor(unit_fc[data_row % 2])
                cell.set_text_props(fontweight="bold", ha="center", va="center", color="#1A237E")
            else:
                cell.set_facecolor(row_fc[data_row % 2])
                cell.set_text_props(ha="center", va="center", color="#212121")

    fig.savefig(path, dpi=300, bbox_inches="tight", facecolor="white", pad_inches=0.25)
    plt.close(fig)


def export_jpg(spec: dict, path: Path):
    if spec["num"] == 4:
        export_rubric_image(spec, path.with_suffix(".png"))
        return

    plt.rcParams.update({
        "font.family": "Malgun Gothic",
        "axes.unicode_minus": False,
    })

    num = spec["num"]
    title = spec["title"]
    headers = spec["headers"]
    rows = spec["rows"]
    n_cols = len(headers)

    # 표 크기에 따른 figure 크기
    n_rows = len(rows)
    fig_w = max(10, n_cols * 2.8)
    fig_h = max(3, 1.2 + n_rows * 0.55)
    if num == 5:
        fig_h = 10
    if num == 4:
        fig_w = 14
        fig_h = 5

    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.axis("off")
    ax.set_title(f"<표 {num}> {title}", fontsize=13, fontweight="bold", pad=16)

    # 셀 텍스트 (줄바꿈)
    wrap_limits = {1: 12, 2: 28, 3: 22, 4: 14, 5: 22, 6: 28, 7: 55}
    limit = wrap_limits.get(num, 20)
    cell_text = []
    for row in rows:
        cell_text.append([_wrap_text(_sanitize_for_jpg(str(c)), limit) for c in row])

    col_labels = [_wrap_text(_sanitize_for_jpg(h), limit) for h in headers]

    tbl = ax.table(
        cellText=cell_text,
        colLabels=col_labels,
        loc="center",
        cellLoc="left",
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(8 if num in (4, 5) else 9)
    tbl.scale(1.0, 1.6 if num == 5 else 1.8)

    # 헤더 스타일
    for j in range(n_cols):
        cell = tbl[0, j]
        cell.set_facecolor("#E8EEF7")
        cell.set_text_props(fontweight="bold", ha="center")

    # 표 7 영역 열 — 병합 시각화 (첫 행만 굵게)
    if num == 7:
        merge_ranges = [(0, 1), (2, 3), (4, 5)]
        for start, end in merge_ranges:
            tbl[start + 1, 0].set_text_props(fontweight="bold")
            for r in range(start + 2, end + 2):
                tbl[r, 0].set_facecolor("#FAFAFA")
                tbl[r, 0].get_text().set_text("")

    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white", pad_inches=0.3)
    plt.close(fig)


def main():
    OUT_DIR.mkdir(exist_ok=True)
    DOWNLOAD_DIR.mkdir(exist_ok=True)
    print(f"Output → {OUT_DIR}")
    print(f"Download → {DOWNLOAD_DIR}\n")

    download_names = {
        4: "표_05_NCS_능력단위별_평가_루브릭",  # 연구계획서 HWP 표 번호
    }

    # 개별 DOCX + JPG/PNG
    for spec in TABLES:
        name = slug(spec["num"], spec["title"])
        docx_path = OUT_DIR / f"{name}.docx"
        jpg_path = OUT_DIR / f"{name}.jpg"
        export_docx(spec, docx_path)
        export_jpg(spec, jpg_path)
        print(f"  [OK] {docx_path.name}")
        if spec["num"] == 4:
            png_path = OUT_DIR / f"{name}.png"
            print(f"  [OK] {png_path.name}")
        else:
            print(f"  [OK] {jpg_path.name}")
        if spec["num"] in download_names:
            dl_png = DOWNLOAD_DIR / f"{download_names[spec['num']]}.png"
            export_rubric_image(spec, dl_png)
            print(f"  [OK] download/{dl_png.name}")

    # 통합 DOCX
    combined = setup_doc()
    cover = combined.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run("연구계획서 표 1~7")
    set_run_font(run, size_pt=16, bold=True)
    combined.add_paragraph()

    for spec in TABLES:
        add_table_to_doc(combined, spec)
        if spec["num"] < 7:
            combined.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    combined_path = OUT_DIR / "표_전체_1-7.docx"
    combined.save(combined_path)
    print(f"\n  [OK] {combined_path.name} (7개 표 통합)")
    print(f"\nDone - {len(TABLES)} tables x (docx + jpg) + 1 combined docx")


if __name__ == "__main__":
    main()
